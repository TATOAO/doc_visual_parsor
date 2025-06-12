"""
Document-native Layout Detection Module

This module provides a document-native implementation of layout detection that works
directly with document structures (like .docx files) using their internal XML and
style information, without converting to images first.
"""

import logging
import tempfile
import os
from pathlib import Path
from typing import Union, List, Dict, Any, Optional
import xml.etree.ElementTree as ET
from zipfile import ZipFile
import re

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise ImportError("Please install python-docx: pip install python-docx")

from ..base.base_detector import BaseSectionDetector
from models.schemas.layout_schemas import (
    StyleInfo, 
    FontInfo, 
    ParagraphFormat, 
    RunInfo, 
    TextAlignment,
    LayoutDetectionResult,
    LayoutElement,
    BoundingBox,
    ElementType
)
from models.schemas.schemas import Section, Positions, DocumentType

logger = logging.getLogger(__name__)

import numpy as np
from PIL import Image
from typing import BinaryIO

InputDataType = Union[
    str,                    # File path as string
    Path,                   # File path as Path object
    bytes,                  # Raw file content
    np.ndarray,            # Image as numpy array
    Image.Image,           # PIL Image object
    BinaryIO,              # File-like object with read() method
    Any                    # For objects with getvalue() method (uploaded files)
]

class DocxLayoutDetector(BaseSectionDetector):
    """
    Document-native Layout Detector for .docx files.
    
    This detector works directly with the document's internal structure,
    analyzing XML elements, styles, and formatting to identify layout elements
    without converting to images.
    """

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.detector = self._initialize_detector()
    
    def _initialize_detector(self):
        """Initialize the detector (no special initialization needed for docx parsing)."""
        pass
    
    def _detect_layout(self, 
                      input_data: InputDataType,
                      confidence_threshold: Optional[float] = None,
                      **kwargs) -> LayoutDetectionResult:
        """
        Core detection method - analyzes document structure to detect layout elements.
        
        Args:
            input_data: Input data (file path, bytes, etc.)
            confidence_threshold: Override default confidence threshold
            **kwargs: Additional detection parameters
            
        Returns:
            LayoutDetectionResult containing detected elements
        """
        try:
            # Load the document using the same logic as generate_section_tree
            if isinstance(input_data, (str, Path)):
                doc = Document(input_data)
            elif hasattr(input_data, 'read'):
                doc = Document(input_data)
            elif hasattr(input_data, 'getvalue'):
                doc = Document(input_data.getvalue())
            else:
                raise ValueError(f"Unsupported input data type: {type(input_data)}")
            
            # Analyze document structure
            sections_data = self._analyze_document_structure(doc)
            
            # Convert sections data to layout elements
            elements = self._create_layout_elements(sections_data, confidence_threshold)
            
            # Create metadata
            metadata = {
                'detection_method': 'document_native',
                'total_paragraphs': len(sections_data),
                'document_type': 'docx'
            }
            
            return LayoutDetectionResult(elements=elements, metadata=metadata)
            
        except Exception as e:
            logger.error(f"Error in document layout detection: {str(e)}")
            # Return empty result on error
            return LayoutDetectionResult(elements=[], metadata={"error": str(e)})
    
    def _create_layout_elements(self, sections_data: List[Dict[str, Any]], 
                               confidence_threshold: Optional[float] = None) -> List['LayoutElement']:
        """
        Convert sections data to layout elements.
        
        Args:
            sections_data: List of section data dictionaries from document analysis
            confidence_threshold: Minimum confidence threshold to apply
            
        Returns:
            List of LayoutElement objects
        """
        elements = []
        element_id = 0
        
        # Set default confidence threshold
        min_confidence = confidence_threshold or 0.8
        
        for data in sections_data:
            # Determine element type based on heading level and style
            element_type = self._map_to_element_type(data)
            
            # Calculate confidence based on style characteristics
            confidence = self._calculate_confidence(data)
            
            # Skip elements below confidence threshold
            if confidence < min_confidence:
                continue
            
            # Create bounding box (approximate, since docx doesn't have pixel coordinates)
            # We use paragraph index and character positions as approximations
            bbox = BoundingBox(
                x1=0.0,  # Left margin
                y1=float(data['paragraph_index'] * 20),  # Approximate line height
                x2=500.0,  # Approximate page width
                y2=float(data['paragraph_index'] * 20 + 15)  # Approximate element height
            )
            
            # Create metadata
            metadata = {
                'paragraph_index': data['paragraph_index'],
                'char_start': data['char_start'],
                'char_end': data['char_end'],
                'heading_level': data['heading_level'],
                'is_heading': data['is_heading'],
                'style_name': data['style_info'].style_name if data['style_info'] else None,
                'detection_method': 'document_native'
            }
            
            # Add font information if available
            if data['style_info'] and data['style_info'].primary_font:
                font = data['style_info'].primary_font
                metadata.update({
                    'font_name': font.name,
                    'font_size': font.size,
                    'font_bold': font.bold,
                    'font_italic': font.italic,
                    'font_underline': font.underline
                })
            
            # Create layout element
            element = LayoutElement(
                id=element_id,
                element_type=element_type,
                confidence=confidence,
                bbox=bbox,
                text=data['text'],
                style=data['style_info'],
                metadata=metadata
            )
            
            elements.append(element)
            element_id += 1
        
        return elements
    
    def _map_to_element_type(self, data: Dict[str, Any]) -> 'ElementType':
        """
        Map section data to appropriate ElementType.
        
        Args:
            data: Section data dictionary
            
        Returns:
            Appropriate ElementType
        """
        if data['is_heading']:
            if data['heading_level'] == 1:
                return ElementType.TITLE
            else:
                return ElementType.HEADING
        else:
            # Check if it might be other types based on content/style
            text = data['text'].lower()
            
            # Simple heuristics for detecting other element types
            if 'table' in text or 'figure' in text:
                if 'table' in text:
                    return ElementType.TABLE_CAPTION
                else:
                    return ElementType.FIGURE_CAPTION
            elif len(data['text']) > 500:  # Long text is likely a paragraph
                return ElementType.PARAGRAPH
            else:
                return ElementType.TEXT
    
    def _calculate_confidence(self, data: Dict[str, Any]) -> float:
        """
        Calculate confidence score for an element based on its characteristics.
        
        Args:
            data: Section data dictionary
            
        Returns:
            Confidence score between 0.0 and 1.0
        """
        confidence = 0.5  # Base confidence
        
        # Higher confidence for headings with clear style indicators
        if data['is_heading']:
            confidence += 0.3
            
            # Even higher confidence for built-in heading styles
            if data['style_info'] and data['style_info'].style_name:
                style_name = data['style_info'].style_name.lower()
                if 'heading' in style_name or 'title' in style_name:
                    confidence += 0.2
        
        # Higher confidence for elements with clear formatting
        if data['style_info'] and data['style_info'].primary_font:
            font = data['style_info'].primary_font
            
            # Bold text gets higher confidence
            if font.bold:
                confidence += 0.1
            
            # Large font size gets higher confidence
            if font.size and font.size >= 14:
                confidence += 0.1
        
        # Text length also affects confidence
        if len(data['text']) > 20:  # Reasonable text length
            confidence += 0.1
        
        return min(confidence, 1.0)  # Cap at 1.0
    
    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported input formats.
        
        Returns:
            List of supported file extensions
        """
        return ['.docx', '.doc']
    
    def get_detector_info(self) -> Dict[str, Any]:
        """
        Get information about the detector.
        
        Returns:
            Dictionary containing detector information
        """
        return {
            'name': 'SectionLayoutDetector',
            'version': '1.0.0',
            'description': 'Document-native layout detector for DOCX files using python-docx',
            'supported_formats': self.get_supported_formats(),
            'detection_method': 'style-based',
            'requires_conversion': False
        }
    
    
    def _extract_style_info(self, paragraph) -> StyleInfo:
        """
        Extract comprehensive style information from a paragraph.
        
        Args:
            paragraph: Document paragraph
            
        Returns:
            StyleInfo object with formatting details
        """
        # Extract paragraph format
        para_format = None
        if paragraph.paragraph_format:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: TextAlignment.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER: TextAlignment.CENTER,
                WD_ALIGN_PARAGRAPH.RIGHT: TextAlignment.RIGHT,
                WD_ALIGN_PARAGRAPH.JUSTIFY: TextAlignment.JUSTIFY,
                WD_ALIGN_PARAGRAPH.DISTRIBUTE: TextAlignment.DISTRIBUTE
            }
            
            para_format = ParagraphFormat(
                alignment=alignment_map.get(paragraph.paragraph_format.alignment, TextAlignment.UNKNOWN),
                left_indent=self._points_to_float(paragraph.paragraph_format.left_indent),
                right_indent=self._points_to_float(paragraph.paragraph_format.right_indent),
                first_line_indent=self._points_to_float(paragraph.paragraph_format.first_line_indent),
                space_before=self._points_to_float(paragraph.paragraph_format.space_before),
                space_after=self._points_to_float(paragraph.paragraph_format.space_after),
                keep_together=paragraph.paragraph_format.keep_together,
                keep_with_next=paragraph.paragraph_format.keep_with_next,
                page_break_before=paragraph.paragraph_format.page_break_before,
                widow_control=paragraph.paragraph_format.widow_control
            )
        
        # Extract run information
        runs = []
        for run in paragraph.runs:
            if run.text:
                # Safely extract font properties
                font_info = FontInfo(
                    name=getattr(run.font, 'name', None),
                    size=self._points_to_float(getattr(run.font, 'size', None)),
                    bold=self._safe_bool_property(run.font, 'bold'),
                    italic=self._safe_bool_property(run.font, 'italic'),
                    underline=self._safe_bool_property(run.font, 'underline')
                )
                
                runs.append(RunInfo(
                    text=run.text,
                    font=font_info
                ))
        
        # Extract primary font (from first run or default)
        primary_font = None
        if runs:
            primary_font = runs[0].font
        
        return StyleInfo(
            style_name=paragraph.style.name if paragraph.style else None,
            style_type='paragraph',
            paragraph_format=para_format,
            runs=runs,
            primary_font=primary_font
        )
    
    def _points_to_float(self, points_value) -> Optional[float]:
        """Convert docx points value to float."""
        if points_value is None:
            return None
        try:
            return float(points_value.pt) if hasattr(points_value, 'pt') else float(points_value)
        except (AttributeError, ValueError):
            return None
    
    def _safe_bool_property(self, font_obj, property_name: str) -> Optional[bool]:
        """Safely extract boolean font properties."""
        try:
            value = getattr(font_obj, property_name, None)
            if value is None:
                return None
            # Handle boolean values directly
            if isinstance(value, bool):
                return value
            # Handle special docx objects that might have boolean interpretation
            if hasattr(value, '__bool__'):
                return bool(value)
            # For underline specifically, check if it's truthy (non-None, non-False)
            if property_name == 'underline':
                return value is not None and value is not False
            return None
        except Exception:
            return None
    
    def _determine_heading_level(self, paragraph, style_info: StyleInfo) -> int:
        """
        Determine if paragraph is a heading and its level.
        
        Args:
            paragraph: Document paragraph
            style_info: Style information
            
        Returns:
            Heading level (0 for non-headings, 1-9 for headings)
        """
        # Check built-in heading styles
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            
            # Built-in heading styles
            heading_patterns = [
                r'heading\s*(\d+)',
                r'title',
                r'subtitle'
            ]
            
            for pattern in heading_patterns:
                match = re.search(pattern, style_name)
                if match:
                    if pattern == r'title':
                        return 1
                    elif pattern == r'subtitle':
                        return 2
                    else:
                        level = int(match.group(1))
                        return min(level, 9)  # Cap at level 9
        
        # Check font characteristics for potential headings
        if style_info.primary_font:
            font = style_info.primary_font
            
            # Large font size suggests heading
            if font.size and font.size >= 14:
                if font.size >= 18:
                    return 1
                elif font.size >= 16:
                    return 2
                else:
                    return 3
            
            # Bold text might be a heading
            if font.bold and len(paragraph.text.strip()) < 100:
                return 3
        
        return 0  # Not a heading
    
    def _build_section_tree(self, sections_data: List[Dict[str, Any]]) -> Section:
        """
        Build hierarchical section tree from section data.
        
        Args:
            sections_data: List of section data dictionaries
            
        Returns:
            Root section with nested subsections
        """
        if not sections_data:
            return Section(
                title="Empty Document",
                content="",
                level=0,
                title_position=Positions.from_docx(0, 0, 0),
                content_position=Positions.from_docx(0, 0, 0)
            )
        
        # Create root section
        root = Section(
            title="Document",
            content="",
            level=0,
            title_position=Positions.from_docx(0, 0, 0),
            content_position=Positions.from_docx(0, 0, 0)
        )
        
        # Stack to track current section hierarchy
        section_stack = [root]
        current_content = []
        
        for data in sections_data:
            if data['is_heading']:
                # Save accumulated content to current section
                if current_content:
                    current_section = section_stack[-1]
                    content_text = '\n'.join(current_content)
                    current_section.content = content_text
                    current_section.content_parsed = content_text
                    current_content = []
                
                # Create new section
                heading_level = data['heading_level']
                new_section = Section(
                    title=data['text'],
                    content="",
                    level=heading_level,
                    title_position=Positions.from_docx(
                        data['paragraph_index'],
                        data['char_start'],
                        data['char_end']
                    ),
                    content_position=Positions.from_docx(
                        data['paragraph_index'],
                        data['char_end'],
                        data['char_end']
                    )
                )
                new_section.title_parsed = data['text']
                
                # Find appropriate parent in stack
                while len(section_stack) > 1 and section_stack[-1].level >= heading_level:
                    section_stack.pop()
                
                # Add to parent section
                parent = section_stack[-1]
                new_section.parent_section = parent
                parent.sub_sections.append(new_section)
                section_stack.append(new_section)
                
            else:
                # Accumulate content
                current_content.append(data['text'])
        
        # Save final accumulated content
        if current_content:
            current_section = section_stack[-1]
            content_text = '\n'.join(current_content)
            current_section.content = content_text
            current_section.content_parsed = content_text
        
        return root


# unit test
# python -m models.layout_detection.document_detector
if __name__ == "__main__":
    detector = SectionLayoutDetector()
    result = detector._detect_layout(input_data="tests/test_data/1-1 买卖合同（通用版）.docx")
    import json 
    json.dump(result.model_dump(), open("layout_detection_result.json", "w"), indent=2, ensure_ascii=False)

    # result = detector.generate_section_tree(input_data="tests/test_data/1-1 买卖合同（通用版）.docx")
    # from models.naive_llm.helpers.section_token_parsor import remove_circular_references
    # remove_circular_references(result)
    # print(result.model_dump_json(indent=2))