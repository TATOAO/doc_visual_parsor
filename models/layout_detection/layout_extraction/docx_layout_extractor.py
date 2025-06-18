"""
Document-native Layout Detection Module

This module provides a document-native implementation of layout detection that works
directly with document structures (like .docx files) using their internal XML and
style information, without converting to images first.
"""

import logging
import tempfile
import os
from pathlib import Path
from typing import Union, List, Dict, Any, Optional
import xml.etree.ElementTree as ET
from zipfile import ZipFile
import re

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise ImportError("Please install python-docx: pip install python-docx")

from ..base.base_layout_extractor import BaseLayoutExtractor
from models.schemas.layout_schemas import (
    StyleInfo, 
    FontInfo, 
    ParagraphFormat, 
    RunInfo, 
    TextAlignment,
    LayoutExtractionResult,
    LayoutElement,
    BoundingBox,
    ElementType
)
from models.schemas.schemas import Section, Positions, DocumentType

logger = logging.getLogger(__name__)

import numpy as np
from PIL import Image
from typing import BinaryIO

InputDataType = Union[
    str,                    # File path as string
    Path,                   # File path as Path object
    bytes,                  # Raw file content
    np.ndarray,            # Image as numpy array
    Image.Image,           # PIL Image object
    BinaryIO,              # File-like object with read() method
    Any                    # For objects with getvalue() method (uploaded files)
]

class DocxLayoutExtrator(BaseLayoutExtractor):
    """
    Document-native Style Extractor for .docx files.
    
    This extractor works directly with the document's internal structure,
    analyzing XML elements, styles, and formatting to extract raw style information
    without determining layout element types.
    """

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.detector = self._initialize_detector()
    
    def _initialize_detector(self):
        """Initialize the detector (no special initialization needed for docx parsing)."""
        pass
    
    def _detect_layout(self, 
                      input_data: InputDataType,
                      confidence_threshold: Optional[float] = None,
                      **kwargs) -> LayoutExtractionResult:
        """
        Core extraction method - analyzes document structure to extract raw style information.
        
        Args:
            input_data: Input data (file path, bytes, etc.)
            confidence_threshold: Not used in raw extraction
            **kwargs: Additional extraction parameters
            
        Returns:
            LayoutExtractionResult containing raw style elements
        """
        try:
            # Load the document
            if isinstance(input_data, (str, Path)):
                doc = Document(input_data)
            elif hasattr(input_data, 'read'):
                doc = Document(input_data)
            elif hasattr(input_data, 'getvalue'):
                doc = Document(input_data.getvalue())
            else:
                raise ValueError(f"Unsupported input data type: {type(input_data)}")
            
            # Extract raw document structure
            raw_elements = self._extract_raw_document_structure(doc)
            
            # Create metadata
            metadata = {
                'extraction_method': 'document_native',
                'total_paragraphs': len(raw_elements),
                'document_type': 'docx'
            }
            
            return LayoutExtractionResult(elements=raw_elements, metadata=metadata)
            
        except Exception as e:
            logger.error(f"Error in document style extraction: {str(e)}")
            # Return empty result on error
            return LayoutExtractionResult(elements=[], metadata={"error": str(e)})
    
    def _extract_raw_document_structure(self, doc: Document) -> List['LayoutElement']:
        """
        Extract raw document structure with style information only.
        
        Args:
            doc: Document object
            
        Returns:
            List of LayoutElement objects with raw style information
        """
        elements = []
        element_id = 0
        char_position = 0
        paragraph_count = 0
        
        # Process paragraphs first
        for paragraph_index, paragraph in enumerate(doc.paragraphs):
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue
            
            # Extract comprehensive style information
            style_info = self._extract_style_info(paragraph)
            
            # Calculate text positions
            text = paragraph.text
            char_start = char_position
            char_end = char_position + len(text)
            char_position = char_end + 1  # +1 for paragraph break
            
            # Determine if this is a heading based on style only (no classification)
            heading_level = self._extract_heading_level_from_style(paragraph, style_info)
            
            # Create approximate bounding box (docx doesn't have pixel coordinates)
            bbox = BoundingBox(
                x1=0.0,  # Left margin
                y1=float(paragraph_index * 20),  # Approximate line height
                x2=500.0,  # Approximate page width
                y2=float(paragraph_index * 20 + 15)  # Approximate element height
            )
            
            # Create metadata with raw information only
            metadata = {
                'paragraph_index': paragraph_index,
                'char_start': char_start,
                'char_end': char_end,
                'heading_level': heading_level,
                'is_heading': heading_level > 0,
                'style_name': style_info.style_name if style_info else None,
                'extraction_method': 'document_native',
                'raw_extraction': True
            }
            
            # Add font information if available
            if style_info and style_info.primary_font:
                font = style_info.primary_font
                metadata.update({
                    'font_name': font.name,
                    'font_size': font.size,
                    'font_bold': font.bold,
                    'font_italic': font.italic,
                    'font_underline': font.underline
                })
            
            # Create layout element with raw information only
            # Note: We set element_type to TEXT as default, classification will happen later
            element = LayoutElement(
                id=element_id,
                element_type=ElementType.PLAIN_TEXT,  # Default type, classification happens later
                confidence=1.0,  # Raw extraction has full confidence in style info
                bbox=bbox,
                text=text,
                style=style_info,
                metadata=metadata
            )
            
            elements.append(element)
            element_id += 1
            paragraph_count += 1
        
        # Process tables
        for table_index, table in enumerate(doc.tables):
            # Extract table content and structure
            table_text, table_metadata = self._extract_table_content(table, table_index)
            
            if table_text:  # Only add non-empty tables
                # Create approximate bounding box for table
                bbox = BoundingBox(
                    x1=0.0,
                    y1=float((paragraph_count + table_index) * 20),
                    x2=500.0,
                    y2=float((paragraph_count + table_index) * 20 + 50)  # Tables are typically taller
                )
                
                # Create table metadata
                metadata = {
                    'table_index': table_index,
                    'char_start': char_position,
                    'char_end': char_position + len(table_text),
                    'extraction_method': 'document_native',
                    'raw_extraction': True,
                    'element_type_override': 'Table',  # Indicate this should be classified as table
                    **table_metadata
                }
                
                # Create table element
                element = LayoutElement(
                    id=element_id,
                    element_type=ElementType.TABLE,  # Mark as table
                    confidence=1.0,
                    bbox=bbox,
                    text=table_text,
                    style=self._extract_table_style_info(table),  # Extract table style information
                    metadata=metadata
                )
                
                elements.append(element)
                element_id += 1
                char_position += len(table_text) + 1
        
        return elements
    
    def _extract_heading_level_from_style(self, paragraph, style_info: StyleInfo) -> int:
        """
        Extract heading level information from paragraph style (raw style info only).
        
        Args:
            paragraph: Document paragraph
            style_info: Style information
            
        Returns:
            Heading level (0 for non-headings, 1-9 for headings) based on style only
        """
        # Check built-in heading styles
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            
            # Built-in heading styles
            heading_patterns = [
                r'heading\s*(\d+)',
                r'title',
                r'subtitle'
            ]
            
            for pattern in heading_patterns:
                match = re.search(pattern, style_name)
                if match:
                    if pattern == r'title':
                        return 1
                    elif pattern == r'subtitle':
                        return 2
                    else:
                        level = int(match.group(1))
                        return min(level, 9)  # Cap at level 9
        
        return 0  # Not a heading based on style
    
    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported input formats.
        
        Returns:
            List of supported file extensions
        """
        return ['.docx', '.doc']
    
    def get_detector_info(self) -> Dict[str, Any]:
        """
        Get information about the extractor.
        
        Returns:
            Dictionary containing extractor information
        """
        return {
            'name': 'DocxStyleExtractor',
            'version': '1.0.0',
            'description': 'Document-native style extractor for DOCX files using python-docx',
            'supported_formats': self.get_supported_formats(),
            'extraction_method': 'style-based',
            'requires_conversion': False,
            'raw_extraction_only': True
        }
    
    def _extract_style_info(self, paragraph) -> StyleInfo:
        """
        Extract comprehensive style information from a paragraph.
        
        Args:
            paragraph: Document paragraph
            
        Returns:
            StyleInfo object with formatting details
        """
        # Extract paragraph format
        para_format = None
        if paragraph.paragraph_format:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: TextAlignment.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER: TextAlignment.CENTER,
                WD_ALIGN_PARAGRAPH.RIGHT: TextAlignment.RIGHT,
                WD_ALIGN_PARAGRAPH.JUSTIFY: TextAlignment.JUSTIFY,
                WD_ALIGN_PARAGRAPH.DISTRIBUTE: TextAlignment.DISTRIBUTE
            }
            
            para_format = ParagraphFormat(
                alignment=alignment_map.get(paragraph.paragraph_format.alignment, TextAlignment.UNKNOWN),
                left_indent=self._points_to_float(paragraph.paragraph_format.left_indent),
                right_indent=self._points_to_float(paragraph.paragraph_format.right_indent),
                first_line_indent=self._points_to_float(paragraph.paragraph_format.first_line_indent),
                space_before=self._points_to_float(paragraph.paragraph_format.space_before),
                space_after=self._points_to_float(paragraph.paragraph_format.space_after),
                keep_together=paragraph.paragraph_format.keep_together,
                keep_with_next=paragraph.paragraph_format.keep_with_next,
                page_break_before=paragraph.paragraph_format.page_break_before,
                widow_control=paragraph.paragraph_format.widow_control
            )
        
        # Extract run information
        runs = []
        for run in paragraph.runs:
            if run.text:
                # Safely extract font properties
                font_info = FontInfo(
                    name=getattr(run.font, 'name', None),
                    size=self._points_to_float(getattr(run.font, 'size', None)),
                    bold=self._safe_bool_property(run.font, 'bold'),
                    italic=self._safe_bool_property(run.font, 'italic'),
                    underline=self._safe_bool_property(run.font, 'underline')
                )
                
                runs.append(RunInfo(
                    text=run.text,
                    font=font_info
                ))
        
        # Extract primary font (from first run or default)
        primary_font = None
        if runs:
            primary_font = runs[0].font
        
        return StyleInfo(
            style_name=paragraph.style.name if paragraph.style else None,
            style_type='paragraph',
            paragraph_format=para_format,
            runs=runs,
            primary_font=primary_font
        )
    
    def _points_to_float(self, points_value) -> Optional[float]:
        """Convert docx points value to float."""
        if points_value is None:
            return None
        try:
            return float(points_value.pt) if hasattr(points_value, 'pt') else float(points_value)
        except (AttributeError, ValueError):
            return None
    
    def _safe_bool_property(self, font_obj, property_name: str) -> Optional[bool]:
        """Safely extract boolean font properties."""
        try:
            value = getattr(font_obj, property_name, None)
            if value is None:
                return None
            # Handle boolean values directly
            if isinstance(value, bool):
                return value
            # Handle special docx objects that might have boolean interpretation
            if hasattr(value, '__bool__'):
                return bool(value)
            # For underline specifically, check if it's truthy (non-None, non-False)
            if property_name == 'underline':
                return value is not None and value is not False
            return None
        except Exception:
            return None
    
    def _build_section_tree(self, sections_data: List[Dict[str, Any]]) -> Section:
        """
        Build hierarchical section tree from section data.
        
        Args:
            sections_data: List of section data dictionaries
            
        Returns:
            Root section with nested subsections
        """
        if not sections_data:
            return Section(
                title="Empty Document",
                content="",
                level=0,
                title_position=Positions.from_docx(0, 0, 0),
                content_position=Positions.from_docx(0, 0, 0)
            )
        
        # Create root section
        root = Section(
            title="Document",
            content="",
            level=0,
            title_position=Positions.from_docx(0, 0, 0),
            content_position=Positions.from_docx(0, 0, 0)
        )
        
        # Stack to track current section hierarchy
        section_stack = [root]
        current_content = []
        
        for data in sections_data:
            if data['is_heading']:
                # Save accumulated content to current section
                if current_content:
                    current_section = section_stack[-1]
                    content_text = '\n'.join(current_content)
                    current_section.content = content_text
                    current_section.content_parsed = content_text
                    current_content = []
                
                # Create new section
                heading_level = data['heading_level']
                new_section = Section(
                    title=data['text'],
                    content="",
                    level=heading_level,
                    title_position=Positions.from_docx(
                        data['paragraph_index'],
                        data['char_start'],
                        data['char_end']
                    ),
                    content_position=Positions.from_docx(
                        data['paragraph_index'],
                        data['char_end'],
                        data['char_end']
                    )
                )
                new_section.title_parsed = data['text']
                
                # Find appropriate parent in stack
                while len(section_stack) > 1 and section_stack[-1].level >= heading_level:
                    section_stack.pop()
                
                # Add to parent section
                parent = section_stack[-1]
                new_section.parent_section = parent
                parent.sub_sections.append(new_section)
                section_stack.append(new_section)
                
            else:
                # Accumulate content
                current_content.append(data['text'])
        
        # Save final accumulated content
        if current_content:
            current_section = section_stack[-1]
            content_text = '\n'.join(current_content)
            current_section.content = content_text
            current_section.content_parsed = content_text
        
        return root

    def _extract_table_content(self, table, table_index: int) -> tuple[str, dict]:
        """
        Extract content and metadata from a table.
        
        Args:
            table: Document table object
            table_index: Index of the table in the document
            
        Returns:
            Tuple of (table_text, table_metadata)
        """
        rows_data = []
        total_cells = 0
        
        for row_index, row in enumerate(table.rows):
            row_data = []
            for cell_index, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_data.append(cell_text)
                total_cells += 1
            rows_data.append(row_data)
        
        # Convert table to text representation
        if rows_data:
            # Create a simple text representation of the table
            table_lines = []
            for row in rows_data:
                # Join cells with tab separator for basic table format
                table_lines.append('\t'.join(row))
            table_text = '\n'.join(table_lines)
        else:
            table_text = ""
        
        # Create table metadata
        table_metadata = {
            'table_rows': len(rows_data),
            'table_cols': len(rows_data[0]) if rows_data else 0,
            'total_cells': total_cells,
            'table_structure': rows_data  # Include the actual table structure
        }
        
        return table_text, table_metadata

    def _extract_table_style_info(self, table) -> StyleInfo:
        """
        Extract comprehensive style information from a table.
        
        Args:
            table: Document table object
            
        Returns:
            StyleInfo object with table formatting details
        """
        # Extract table-level style information
        table_style_info = {}
        
        # Get table style name if available
        style_name = None
        if hasattr(table, 'style') and table.style:
            style_name = table.style.name
            table_style_info['style_name'] = style_name
            
        # Extract table alignment if available
        if hasattr(table, 'alignment') and table.alignment:
            table_style_info['alignment'] = str(table.alignment)
            
        # Extract table width information
        if hasattr(table, 'autofit'):
            table_style_info['autofit'] = table.autofit
            
        # Extract font information from first cell with text
        primary_font = None
        runs = []
        
        # Iterate through table cells to find text and extract font information
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():  # Only process paragraphs with text
                        for run in paragraph.runs:
                            if run.text:
                                # Extract font information similar to paragraph runs
                                font_info = FontInfo(
                                    name=getattr(run.font, 'name', None),
                                    size=self._points_to_float(getattr(run.font, 'size', None)),
                                    bold=self._safe_bool_property(run.font, 'bold'),
                                    italic=self._safe_bool_property(run.font, 'italic'),
                                    underline=self._safe_bool_property(run.font, 'underline')
                                )
                                
                                runs.append(RunInfo(
                                    text=run.text,
                                    font=font_info
                                ))
                                
                                # Set primary font from first run found
                                if primary_font is None:
                                    primary_font = font_info
                        
                        # Exit early after finding first text content
                        if runs:
                            break
                if runs:
                    break
            if runs:
                break
        
        return StyleInfo(
            style_name=style_name,
            style_type='table',
            paragraph_format=None,  # Tables don't have paragraph-level formatting
            runs=runs,
            primary_font=primary_font,
            table_style=table_style_info if table_style_info else None,
            custom_properties={'element_type': 'table'}
        )


# unit test
# python -m models.layout_detection.layout_extraction.docx_layout_extractor
if __name__ == "__main__":
    detector = DocxLayoutExtrator()
    result = detector._detect_layout(input_data="tests/test_data/1-1 买卖合同（通用版）.docx")
    import json 

    # remove all runs
    # for element in result.elements:
    #     if element.style and element.style.runs:
    #         element.style.runs = [] 

    json.dump(result.model_dump(), open("layout_detection_result_with_runs.json", "w"), indent=2, ensure_ascii=False)
