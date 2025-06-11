"""
Document-native Layout Detection Module

This module provides a document-native implementation of layout detection that works
directly with document structures (like .docx files) using their internal XML and
style information, without converting to images first.
"""

import logging
import tempfile
import os
from pathlib import Path
from typing import Union, List, Dict, Any, Optional
import xml.etree.ElementTree as ET
from zipfile import ZipFile
import re

try:
    import docx
    from docx.document import Document
except ImportError:
    raise ImportError("Please install python-docx: pip install python-docx")

from .base_detector import (
    BaseLayoutDetector, 
    LayoutDetectionResult, 
    LayoutElement, 
    BoundingBox, 
    ElementType,
    StyleInfo,
    FontInfo,
    ParagraphFormat,
    RunInfo,
    TextAlignment
)

logger = logging.getLogger(__name__)

class DocumentLayoutDetector(BaseLayoutDetector):
    """
    Document-native Layout Detector for .docx files.
    
    This detector works directly with the document's internal structure,
    analyzing XML elements, styles, and formatting to identify layout elements
    without converting to images.
    """
    
    def __init__(self, 
                 confidence_threshold: float = 0.9,  # High confidence since it's rule-based
                 device: str = "cpu",  # CPU-based processing
                 extract_text: bool = True,
                 analyze_styles: bool = True,
                 **kwargs):
        """
        Initialize the document-native layout detector.
        
        Args:
            confidence_threshold: Minimum confidence for detections (rule-based, so high)
            device: Device to use (always CPU for document processing)
            extract_text: Whether to extract text content
            analyze_styles: Whether to analyze document styles
            **kwargs: Additional parameters
        """
        super().__init__(confidence_threshold=confidence_threshold, device="cpu", **kwargs)
        
        self.extract_text = extract_text
        self.analyze_styles = analyze_styles
        self.document = None
        self.styles_map = {}
        
    def _initialize_detector(self) -> None:
        """Initialize the detector (minimal setup for document processing)."""
        logger.info("Document-native layout detector initialized")
    
    def _detect_layout(self, 
                      input_data: Any,
                      confidence_threshold: Optional[float] = None,
                      **kwargs) -> LayoutDetectionResult:
        """
        Core document-native layout detection method.
        
        Args:
            input_data: Input document (file path, bytes, or file-like object)
            confidence_threshold: Override default confidence threshold
            **kwargs: Additional detection parameters
            
        Returns:
            LayoutDetectionResult containing detected elements
        """
        # Load document
        doc = self._load_document(input_data)
        if doc is None:
            return LayoutDetectionResult([])
        
        # Extract layout elements
        elements = []
        element_id = 0
        
        # Analyze document structure
        if self.analyze_styles:
            self._analyze_document_styles(doc)
        
        # Process paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
                
            element_type, confidence = self._classify_paragraph(paragraph, para_idx)
            
            # Create layout element
            element = LayoutElement(
                id=element_id,
                element_type=element_type,
                confidence=confidence,
                text=paragraph.text.strip() if self.extract_text else None,
                style=self._extract_paragraph_style(paragraph),
                metadata={
                    'paragraph_index': para_idx,
                    'detection_method': 'document_native',
                    'style_name': paragraph.style.name if paragraph.style else 'Normal'
                }
            )
            
            elements.append(element)
            element_id += 1
        
        # Process tables
        for table_idx, table in enumerate(doc.tables):
            element = LayoutElement(
                id=element_id,
                element_type=ElementType.TABLE,
                confidence=1.0,  # High confidence for explicit table elements
                text=self._extract_table_text(table) if self.extract_text else None,
                metadata={
                    'table_index': table_idx,
                    'detection_method': 'document_native',
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0
                }
            )
            
            elements.append(element)
            element_id += 1
        
        # Process headers and footers
        header_footer_elements = self._extract_header_footer_elements(doc)
        for hf_element in header_footer_elements:
            hf_element.id = element_id
            elements.append(hf_element)
            element_id += 1
        
        return LayoutDetectionResult(elements)
    
    def _load_document(self, input_data: Any) -> Optional[Document]:
        """Load document from various input formats."""
        try:
            # Save uploaded file to temporary location if needed
            if isinstance(input_data, bytes):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                    tmp_file.write(input_data)
                    tmp_file_path = tmp_file.name
                doc = docx.Document(tmp_file_path)
                os.unlink(tmp_file_path)  # Clean up
                return doc
            elif isinstance(input_data, (str, Path)):
                return docx.Document(str(input_data))
            elif hasattr(input_data, 'getvalue'):
                # Handle uploaded file objects
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                    tmp_file.write(input_data.getvalue())
                    tmp_file_path = tmp_file.name
                doc = docx.Document(tmp_file_path)
                os.unlink(tmp_file_path)  # Clean up
                return doc
            else:
                logger.error(f"Unsupported input type: {type(input_data)}")
                return None
                
        except Exception as e:
            logger.error(f"Failed to load document: {str(e)}")
            return None
    
    def _analyze_document_styles(self, doc: Document) -> None:
        """Analyze document styles to improve classification."""
        self.styles_map = {}
        
        try:
            for style in doc.styles:
                style_name = style.name.lower()
                self.styles_map[style_name] = {
                    'name': style.name,
                    'type': str(style.type),
                    'builtin': style.builtin
                }
        except Exception as e:
            logger.warning(f"Could not analyze styles: {str(e)}")
    
    def _classify_paragraph(self, paragraph, para_idx: int) -> tuple[ElementType, float]:
        """Classify paragraph based on style and content."""
        style_name = paragraph.style.name.lower() if paragraph.style else 'normal'
        text = paragraph.text.strip().lower()
        
        # High confidence classifications based on style
        if 'heading' in style_name:
            if 'heading 1' in style_name:
                return ElementType.TITLE, 1.0
            else:
                return ElementType.HEADING, 1.0
        
        if 'title' in style_name:
            return ElementType.TITLE, 1.0
        
        if 'caption' in style_name:
            if 'table' in style_name:
                return ElementType.TABLE_CAPTION, 1.0
            else:
                return ElementType.FIGURE_CAPTION, 1.0
        
        if 'header' in style_name:
            return ElementType.HEADER, 1.0
        
        if 'footer' in style_name:
            return ElementType.FOOTER, 1.0
        
        # Content-based classification
        if para_idx == 0 and len(text) < 100:
            return ElementType.TITLE, 0.8
        
        # Check for references
        if any(ref_word in text for ref_word in ['references', 'bibliography', 'works cited']):
            return ElementType.REFERENCE, 0.9
        
        # Check for equations (look for mathematical symbols)
        if re.search(r'[∑∏∫∂∆∇±≤≥≠∞√]|\\[a-zA-Z]+|\$.*\$', text):
            return ElementType.EQUATION, 0.8
        
        # Check for lists
        if (text.startswith(('•', '-', '*')) or 
            re.match(r'^\d+[\.\)]\s', text) or
            re.match(r'^[a-zA-Z][\.\)]\s', text)):
            return ElementType.LIST, 0.9
        
        # Default to text/paragraph
        if len(text) > 20:
            return ElementType.PARAGRAPH, 0.7
        else:
            return ElementType.TEXT, 0.7
    
    def _extract_paragraph_style(self, paragraph) -> StyleInfo:
        """Extract style information from paragraph."""
        try:
            # Basic style information
            style_name = paragraph.style.name if paragraph.style else None
            style_type = str(paragraph.style.type) if paragraph.style else None
            builtin = paragraph.style.builtin if paragraph.style else None
            
            # Extract paragraph formatting
            paragraph_format = None
            if paragraph.paragraph_format:
                pf = paragraph.paragraph_format
                
                # Convert alignment to our enum
                alignment = None
                if pf.alignment:
                    alignment_str = str(pf.alignment).lower()
                    if 'left' in alignment_str:
                        alignment = TextAlignment.LEFT
                    elif 'center' in alignment_str:
                        alignment = TextAlignment.CENTER
                    elif 'right' in alignment_str:
                        alignment = TextAlignment.RIGHT
                    elif 'justify' in alignment_str:
                        alignment = TextAlignment.JUSTIFY
                    elif 'distribute' in alignment_str:
                        alignment = TextAlignment.DISTRIBUTE
                    else:
                        alignment = TextAlignment.UNKNOWN
                
                paragraph_format = ParagraphFormat(
                    alignment=alignment,
                    left_indent=pf.left_indent,
                    right_indent=pf.right_indent,
                    first_line_indent=pf.first_line_indent,
                    space_before=pf.space_before,
                    space_after=pf.space_after,
                    line_spacing=pf.line_spacing
                )
            
            # Extract run formatting
            runs = []
            primary_font = None
            
            for i, run in enumerate(paragraph.runs):
                if not run.text.strip():  # Skip empty runs
                    continue
                    
                font_info = FontInfo()
                if run.font:
                    font_info.name = run.font.name
                    font_info.size = float(run.font.size.pt) if run.font.size else None
                    font_info.bold = run.font.bold
                    font_info.italic = run.font.italic
                    font_info.underline = run.font.underline
                    # Note: color extraction would require more complex handling
                
                run_info = RunInfo(
                    text=run.text,
                    font=font_info
                )
                runs.append(run_info)
                
                # Use first run's font as primary font
                if i == 0:
                    primary_font = font_info
            
            return StyleInfo(
                style_name=style_name,
                style_type=style_type,
                builtin=builtin,
                paragraph_format=paragraph_format,
                runs=runs if runs else None,
                primary_font=primary_font
            )
                
        except Exception as e:
            logger.warning(f"Could not extract style info: {str(e)}")
            return StyleInfo()  # Return empty StyleInfo on error
    
    def _extract_table_text(self, table) -> str:
        """Extract text content from table."""
        try:
            text_content = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(' | '.join(row_text))
            return '\n'.join(text_content)
        except Exception as e:
            logger.warning(f"Could not extract table text: {str(e)}")
            return ""
    
    def _extract_header_footer_elements(self, doc: Document) -> List[LayoutElement]:
        """Extract header and footer elements."""
        elements = []
        
        try:
            # Extract headers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            element = LayoutElement(
                                id=0,  # Will be set by caller
                                element_type=ElementType.HEADER,
                                confidence=1.0,
                                text=para.text.strip() if self.extract_text else None,
                                style=self._extract_paragraph_style(para),
                                metadata={
                                    'detection_method': 'document_native',
                                    'section_type': 'header'
                                }
                            )
                            elements.append(element)
                
                # Extract footers
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            element = LayoutElement(
                                id=0,  # Will be set by caller
                                element_type=ElementType.FOOTER,
                                confidence=1.0,
                                text=para.text.strip() if self.extract_text else None,
                                style=self._extract_paragraph_style(para),
                                metadata={
                                    'detection_method': 'document_native',
                                    'section_type': 'footer'
                                }
                            )
                            elements.append(element)
        except Exception as e:
            logger.warning(f"Could not extract header/footer elements: {str(e)}")
        
        return elements
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported input formats."""
        return ['.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
    
    def get_detector_info(self) -> Dict[str, Any]:
        """Get information about the document detector."""
        return {
            'detector_type': 'document_native',
            'supported_formats': self.get_supported_formats(),
            'confidence_threshold': self.confidence_threshold,
            'extract_text': self.extract_text,
            'analyze_styles': self.analyze_styles,
            'capabilities': [
                'heading_detection',
                'table_detection', 
                'header_footer_detection',
                'style_analysis',
                'text_extraction',
                'list_detection',
                'equation_detection'
            ]
        }
    
    def validate_input(self, input_data: Any) -> bool:
        """Validate if input data is supported by document detector."""
        if isinstance(input_data, (str, Path)):
            path = Path(input_data)
            return path.exists() and path.suffix.lower() == '.docx'
        elif isinstance(input_data, bytes):
            # Check if it's likely a DOCX file (ZIP signature)
            return input_data.startswith(b'PK')
        elif hasattr(input_data, 'getvalue'):
            # Handle uploaded file objects
            return True
        else:
            return False
    
    def visualize(self, 
                  input_data: Any,
                  result: Optional[LayoutDetectionResult] = None,
                  save_path: Optional[str] = None,
                  **kwargs) -> None:
        """
        Visualize detection results as text output.
        
        Args:
            input_data: Input document
            result: Detection result (if None, will run detection)
            save_path: Path to save the text output
            **kwargs: Additional visualization parameters
        """
        # Run detection if result not provided
        if result is None:
            result = self.detect(input_data)
        
        # Create text visualization
        output_lines = []
        output_lines.append("Document Layout Analysis Results")
        output_lines.append("=" * 40)
        output_lines.append("")
        
        # Statistics
        stats = result.get_statistics()
        output_lines.append(f"Total Elements: {stats['total_elements']}")
        output_lines.append("Element Types:")
        for elem_type, count in stats['element_types'].items():
            output_lines.append(f"  {elem_type}: {count}")
        output_lines.append("")
        
        # Element details
        output_lines.append("Detected Elements:")
        output_lines.append("-" * 20)
        
        for element in result.get_elements():
            output_lines.append(f"ID: {element.id}")
            output_lines.append(f"Type: {element.element_type.value}")
            output_lines.append(f"Confidence: {element.confidence:.2f}")
            
            if element.text:
                preview_text = element.text[:100] + "..." if len(element.text) > 100 else element.text
                output_lines.append(f"Text: {preview_text}")
            
            if element.metadata:
                output_lines.append(f"Metadata: {element.metadata}")
            
            output_lines.append("")
        
        # Output to file or print
        output_text = "\n".join(output_lines)
        
        if save_path:
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(output_text)
            logger.info(f"Layout analysis saved to: {save_path}")
        else:
            print(output_text) 


# unit test
# python -m models.layout_detection.document_detector
if __name__ == "__main__":
    detector = DocumentLayoutDetector()
    result = detector.visualize(input_data="tests/test_data/1-1 买卖合同（通用版）.docx", save_path="tests/test_data/1-1 买卖合同（通用版）.txt")
    print(result)