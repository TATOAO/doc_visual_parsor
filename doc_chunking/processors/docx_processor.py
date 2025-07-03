import io
from docx import Document
import logging

logger = logging.getLogger(__name__)

def extract_docx_content(uploaded_file):
    """
    Extract text content from a DOCX file.
    
    Args:
        uploaded_file: File-like object with DOCX content
        
    Returns:
        String with extracted text content
    """
    logger.info("Starting DOCX content extraction")
    
    try:
        # Get DOCX content
        if hasattr(uploaded_file, 'getvalue'):
            docx_content = uploaded_file.getvalue()
        elif hasattr(uploaded_file, 'read'):
            docx_content = uploaded_file.read()
        else:
            docx_content = uploaded_file
            
        # Open DOCX from memory
        doc = Document(io.BytesIO(docx_content))
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                text_content.append(paragraph.text)
        
        # Join all paragraphs with newlines
        full_text = "\n".join(text_content)
        
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        return full_text
        
    except Exception as e:
        logger.error(f"Error extracting DOCX content: {str(e)}")
        return ""


def extract_docx_structure(uploaded_file):
    """
    Extract structured content from a DOCX file including headings and paragraphs.
    
    Args:
        uploaded_file: File-like object with DOCX content
        
    Returns:
        List of dictionaries with content structure
    """
    logger.info("Starting DOCX structure extraction")
    
    try:
        # Get DOCX content
        if hasattr(uploaded_file, 'getvalue'):
            docx_content = uploaded_file.getvalue()
        elif hasattr(uploaded_file, 'read'):
            docx_content = uploaded_file.read()
        else:
            docx_content = uploaded_file
            
        # Open DOCX from memory
        doc = Document(io.BytesIO(docx_content))
        
        structured_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                # Check if it's a heading
                style_name = paragraph.style.name.lower()
                is_heading = 'heading' in style_name
                
                # Extract heading level if it's a heading
                heading_level = 0
                if is_heading:
                    try:
                        heading_level = int(style_name.split()[-1])
                    except:
                        heading_level = 1
                
                structured_content.append({
                    'text': paragraph.text,
                    'is_heading': is_heading,
                    'heading_level': heading_level,
                    'style': paragraph.style.name
                })
        
        logger.info(f"Successfully extracted {len(structured_content)} structured elements from DOCX")
        return structured_content
        
    except Exception as e:
        logger.error(f"Error extracting DOCX structure: {str(e)}")
        return [] 